import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def md_to_docx(md_file, docx_file):
    # Leer el contenido del archivo Markdown
    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.read()

    # Crear un nuevo documento Word
    doc = Document()

    # Función para obtener o crear un estilo
    def get_or_create_style(name, font_size, bold=False):
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(font_size)
        style.font.bold = bold
        return style

    # Definir estilos
    style_heading1 = get_or_create_style('Heading 1', 18, bold=True)
    style_heading2 = get_or_create_style('Heading 2', 16, bold=True)
    style_normal = get_or_create_style('Normal', 12)

    # Función para agregar párrafos con el estilo correcto
    def add_paragraph_with_style(text, style):
        p = doc.add_paragraph(text)
        p.style = style
        return p

    # Procesar el contenido
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            add_paragraph_with_style(line[2:], style_heading1)
        elif line.startswith('## '):
            add_paragraph_with_style(line[3:], style_heading2)
        elif line.startswith('* '):
            add_paragraph_with_style('• ' + line[2:], style_normal)
        elif line.strip() != '':
            add_paragraph_with_style(line, style_normal)

    # Guardar el documento
    doc.save(docx_file)
    print(f"Documento Word creado: {docx_file}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Uso: python src/md_to_docx.py <ruta_del_archivo_markdown>")
        sys.exit(1)

    md_file = sys.argv[1]
    if not md_file.endswith('.md'):
        print("El archivo debe tener extensión .md")
        sys.exit(1)

    if not os.path.exists(md_file):
        print(f"El archivo {md_file} no existe.")
        sys.exit(1)

    docx_file = os.path.splitext(md_file)[0] + '.docx'
    md_to_docx(md_file, docx_file)